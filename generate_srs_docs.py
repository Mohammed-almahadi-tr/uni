import os
import sys
import subprocess
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_docx(target_path):
    doc = docx.Document()

    # Configure Margins (0.75 in)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        
        # Configure Header & Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("UniFlow Enterprise ERP — Software Requirements Specification")
        hrun.font.name = "Segoe UI"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(148, 163, 184)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        frun = fp.add_run("Version 2.0.0 | ISO/IEC/IEEE 29148 Standard Compliant | Confidential")
        frun.font.name = "Segoe UI"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(148, 163, 184)

    # Base Colors
    NAVY = RGBColor(30, 58, 138)       # #1E3A8A
    BLUE = RGBColor(37, 99, 235)       # #2563EB
    DARK = RGBColor(31, 41, 55)        # #1F2937
    MUTED = RGBColor(107, 114, 128)    # #6B7280
    BG_HEADER = "1E3A8A"
    BG_ALT = "F8FAFC"

    def add_title(text, subtitle=""):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = "Segoe UI"
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = NAVY

        if subtitle:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(12)
            run2 = p2.add_run(subtitle)
            run2.font.name = "Segoe UI"
            run2.font.size = Pt(13)
            run2.font.color.rgb = BLUE

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Segoe UI"
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = NAVY

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Segoe UI"
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = BLUE

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Segoe UI"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = DARK

    def add_p(text, bold_prefix="", italic=False, space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.font.name = "Segoe UI"
            r_b.font.size = Pt(10)
            r_b.font.bold = True
            r_b.font.color.rgb = DARK
        run = p.add_run(text)
        run.font.name = "Segoe UI"
        run.font.size = Pt(10)
        run.font.italic = italic
        run.font.color.rgb = DARK
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.font.name = "Segoe UI"
            r_b.font.size = Pt(9.5)
            r_b.font.bold = True
            r_b.font.color.rgb = DARK
        run = p.add_run(text)
        run.font.name = "Segoe UI"
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK
        return p

    def add_callout(text, title="CRITICAL REQUIREMENT"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.9)
        set_cell_background(cell, "F0F9FF")
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="24" w:space="0" w:color="0284C7"/>'
            f'<w:top w:val="none"/>'
            f'<w:bottom w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(2)
        r_title = cp.add_run(f"[{title}] ")
        r_title.font.name = "Segoe UI"
        r_title.font.size = Pt(9.5)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(2, 132, 199)
        
        r_text = cp.add_run(text)
        r_text.font.name = "Segoe UI"
        r_text.font.size = Pt(9.5)
        r_text.font.color.rgb = DARK

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_styled_table(headers, rows, col_widths=None):
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, color="E2E8F0", sz="4")

        # Header Row
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(headers):
            cell = hdr_cells[idx]
            if col_widths and idx < len(col_widths):
                cell.width = Inches(col_widths[idx])
            set_cell_background(cell, BG_HEADER)
            set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = "Segoe UI"
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

        # Body Rows
        for r_idx, row_data in enumerate(rows):
            row_cells = table.rows[r_idx + 1].cells
            bg_fill = BG_ALT if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                cell = row_cells[c_idx]
                if col_widths and c_idx < len(col_widths):
                    cell.width = Inches(col_widths[c_idx])
                set_cell_background(cell, bg_fill)
                set_cell_margins(cell, top=90, bottom=90, left=120, right=120)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(str(val))
                run.font.name = "Segoe UI"
                run.font.size = Pt(8.5)
                run.font.color.rgb = DARK

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==================== DOCUMENT CONTENT ====================

    # Title & Metadata Banner
    add_title("Software Requirements Specification (SRS)", "Multi-Tenant University ERP & White-Label Web Platform")
    
    meta_headers = ["Attribute", "Specification Details"]
    meta_rows = [
        ["Project Name", "UniFlow Enterprise ERP (Next-Gen Multi-Tenant Higher-Education Web Platform)"],
        ["Document Version", "2.0.0 (Comprehensive Cloud-Native Architecture)"],
        ["Standard Compliance", "IEEE Std 830-1998 / ISO/IEC/IEEE 29148 Software Requirements Standards"],
        ["Target Audience", "University Executives, System Architects, Full-Stack Engineering, QA, and Product Management"],
        ["Architecture Model", "Cloud-Native SaaS, Multi-Tenant Partitioning, Bilingual (Arabic RTL / English LTR)"],
        ["Security & Data Model", "PostgreSQL RLS, Double-Entry Atomic Ledger, Immutable Hash-Chained Audit Trail, Argon2id"]
    ]
    add_styled_table(meta_headers, meta_rows, [2.0, 4.9])

    # 1. Introduction
    add_h1("1. Introduction")
    add_h2("1.1 Purpose & Executive Overview")
    add_p(
        "This document establishes the authoritative Software Requirements Specification (SRS) for UniFlow Enterprise ERP, "
        "a state-of-the-art, multi-tenant, cloud-native web platform designed for higher education institutions (Universities, Colleges, and Academies). "
        "The system delivers an integrated enterprise suite unifying dynamic white-label public portals, academic management, digital admissions, "
        "student lifecycle management, and a rigorous, double-entry financial accounting engine tailored to institutional operations."
    )
    add_p("The platform is engineered around four core foundational pillars:")
    add_bullet("Dynamic CMS allowing each university to host custom-branded landing pages, faculty explorers, and online admissions portals.", "1. White-Label Public Portal: ")
    add_bullet("Hierarchical structures (Faculties, Degrees, Programs, Batches), student profiles with document checklists, and semester registration engines.", "2. Academic & Student Lifecycle: ")
    add_bullet("5-tier Chart of Accounts, maker-checker voucher approval workflows, fee catalogs, deferred revenue recognition, cashiering, cheques, budgeting, fixed assets, and financial statements.", "3. Comprehensive Financial Engine: ")
    add_bullet("Built with shadcn UI, Tailwind CSS, TypeScript, and React/Next.js, supporting instantaneous zero-reload switching between Arabic (RTL) and English (LTR).", "4. Modern Bilingual Design System: ")

    add_h2("1.2 System Scope & Functional Decomposition")
    add_p(
        "The UniFlow platform provides an end-to-end operational envelope divided into five primary functional modules and an advanced academic records roadmap:"
    )
    add_bullet("Dynamic theme customization, hero banners, public faculty catalogs, online admissions portal, news and academic calendars.", "• Module Cluster A — Public Web & CMS: ")
    add_bullet("Academic hierarchy setup, multi-step student intake, biometric/photo capture, medical examinations, and semester registration engines.", "• Module Cluster B — Academic & Student Operations: ")
    add_bullet("5-tier COA, multi-line journal vouchers, fiscal period controls, multi-channel cashiering, cheque clearing pipeline, budgeting, fixed assets depreciation, and procurement with encumbrance.", "• Module Cluster C — University Financial Engine: ")
    add_bullet("Database-enforced PostgreSQL Row-Level Security (RLS), Role-Based Access Control (RBAC), tamper-evident hash-chained audit trails, and multi-currency exchange rate management.", "• Module Cluster D — Multi-Tenancy & Platform Administration: ")
    add_bullet("Course catalogs, prerequisite trees, timetable scheduling, attendance tracking with 75% rule enforcement, GPA calculation, exam invigilation, and degree audits.", "• Module Cluster E — Academic Records & Grading: ")

    add_h2("1.3 Definitions, Acronyms, and Abbreviations")
    def_headers = ["Term / Acronym", "Definition & Context"]
    def_rows = [
        ["SaaS", "Software as a Service — Multi-tenant cloud-hosted institutional platform."],
        ["COA", "Chart of Accounts (دليل الحسابات) — Standardized 5-tier hierarchical accounting tree."],
        ["GL", "General Ledger (دفتر الأستاذ العام) — Central repository of all financial transactions."],
        ["JV", "Journal Voucher (قيد اليومية) — Multi-line double-entry financial voucher."],
        ["RV / PV", "Receipt Voucher (سند قبض) / Payment Voucher (سند صرف) — Inward/outward monetary documents."],
        ["AR / AP", "Accounts Receivable (حسابات المدينين / الطلاب) / Accounts Payable (حسابات الموردين)."],
        ["RBAC", "Role-Based Access Control — Granular permission assignment per role and institutional department."],
        ["RTL / LTR", "Right-To-Left (Arabic) / Left-To-Right (English) layout rendering."],
        ["Maker-Checker", "Separation of duties workflow requiring voucher creation and approval by distinct authorized actors."],
        ["NBV", "Net Book Value (القيمة الدفترية للأصل) — Historical cost minus accumulated depreciation."],
        ["Encumbrance", "Commitment against a budget line created upon Purchase Order approval prior to disbursement."],
        ["Functional Currency", "The institutional base currency in which financial ledgers and statutory statements are maintained."],
        ["Fiscal Period", "Accounting calendar periods (monthly/quarterly) subject to strict open/closed posting locks."],
        ["Idempotency Key", "Unique client-generated token ensuring that duplicate API submissions create exactly one record."],
        ["تفقيط (Tafqeet)", "Algorithmic translation of monetary numbers into formal written Arabic words on official receipts."]
    ]
    add_styled_table(def_headers, def_rows, [1.8, 5.1])

    # 2. Overall Description
    add_h1("2. Overall Description & System Architecture")
    add_h2("2.1 Product Perspective & Web Architecture")
    add_p(
        "UniFlow is structured as a cloud-native, modern multi-tier web application designed for high concurrency, zero-trust security, and horizontal scalability. "
        "The architecture is partitioned into distinct operational tiers:"
    )
    add_bullet("Next.js / React with TypeScript, Tailwind CSS, and shadcn UI component library. Fully responsive across desktop, tablet, and mobile cashier POS devices.", "1. Presentation Layer (Frontend): ")
    add_bullet("Modular RESTful and GraphQL API services enforcing JWT/session validation, role-based authorization, request payload validation (Zod), and rate limiting.", "2. Application & API Layer: ")
    add_bullet("PostgreSQL with Row-Level Security (RLS) enforcing tenant isolation at the database level, transactional double-entry balancing triggers, and append-only audit tables.", "3. Persistence & Ledger Layer: ")
    add_bullet("Redis for high-speed session management, permission caching, and idempotency token evaluation.", "4. Caching & Idempotency Store: ")
    add_bullet("S3-compatible object storage with server-side AES-256 encryption for student identification documents, certificates, and voucher attachments.", "5. Secure Document Storage: ")

    add_h2("2.2 Core System Capabilities Matrix")
    cap_headers = ["Capability", "Web Platform Implementation", "Architectural & Security Standard"]
    cap_rows = [
        ["Authentication & IAM", "Argon2id password hashing, HTTP-only secure cookies, MFA for financial approvals, RBAC matrix.", "OWASP Top 10 Compliant"],
        ["White-Label CMS", "Dynamic per-tenant theme generator, custom domains, HSL color tokens, typography selection.", "Zero-Code Branding Engine"],
        ["Student Intake", "Multi-step admission wizard, structured 4-part names, biometric/photo capture, document checklist.", "Client & Server Schema Validation"],
        ["Medical Fitness", "Standardized examination panel, blood group, infectious disease screening, fit/unfit clearance.", "Encrypted Medical Records Vault"],
        ["Registration Engine", "Atomic semester registration coupled with immediate double-entry ledger posting in a single transaction.", "ACID Guaranteed Transaction"],
        ["Chart of Accounts", "5-tier normalized tree with parent keys, account codes, normal balance flags, and bilingual titles.", "Strict Level-5 Postability"],
        ["General Ledger", "Immutable transactional ledger (`transaction_headers` / `transaction_lines`), balanced sum constraint.", "Database Trigger Enforced"],
        ["Voucher Approvals", "Maker-Checker multi-stage workflow: Draft → Pending Review → Pending Approval → Posted.", "Immutable State Transition Log"],
        ["Cashiering Terminal", "Multi-channel fee collection (Cash, Bank, Cheque, Online Gateway), automated receipt PDF generation.", "Idempotency Token Enforced"],
        ["Cheque Clearing", "State machine (Received → In Collection → Deposited → Cleared / Bounced) with automatic GL entries.", "Lifecycle Journal Automation"],
        ["Fixed Assets", "Automated straight-line depreciation engine, persisted schedule per asset, period-locked batch runs.", "Idempotent Batch Dispatch"],
        ["Budget & Encumbrance", "Multi-dimensional budget lines with real-time encumbrance tracking fed from Purchase Orders.", "Real-Time Overrun Prevention"],
        ["Financial Reporting", "Incremental balance maintenance (`account_period_balances`), real-time Trial Balance and Balance Sheet.", "< 100ms Aggregate Retrieval"],
        ["Arabic / English Engine", "Native RTL mirroring, font shaping (Cairo/Inter), dual Hijri/Gregorian calendar, Arabic monetary تفقيط.", "Unicode BiDi Compliance"]
    ]
    add_styled_table(cap_headers, cap_rows, [1.5, 3.4, 2.0])

    add_h2("2.3 System Invariants & Architectural Principles")
    add_p(
        "To guarantee institutional-grade data integrity and financial compliance, the UniFlow web platform enforces the following non-negotiable architectural invariants:"
    )
    add_bullet("Every journal voucher must satisfy Sum(Debits) = Sum(Credits) in the tenant's functional currency. This is validated by both client/API schemas and a deferred PostgreSQL database trigger.", "1. Double-Entry Balance: ")
    add_bullet("All monetary values are stored strictly as numeric(19,4) to completely eliminate floating-point rounding errors.", "2. Monetary Precision: ")
    add_bullet("Every state-changing financial submission must include an idempotency key. Duplicate requests within a 24-hour window return the cached response without creating duplicate records.", "3. Financial Idempotency: ")
    add_bullet("Posted financial transactions cannot be updated or deleted by any user role (including Super Admin). Corrections require an explicit, linked reversal voucher with mandatory justification.", "4. Immutability & Reversals: ")
    add_bullet("Every operational table includes a tenant_id column governed by PostgreSQL Row-Level Security policies, ensuring mathematical data isolation across all university clients.", "5. Multi-Tenant Isolation: ")
    add_bullet("All data queries utilize parameterized prepared statements or typed ORM builders. Dynamic SQL concatenation is strictly barred.", "6. Injection Immunity: ")

    add_h2("2.4 User Classes, Personas & Roles (RBAC)")
    add_p("The platform provides pre-configured role archetypes with granular permission matrices:")
    add_bullet("Global platform governance, tenant onboarding, subscription license tier enforcement, and platform-wide monitoring.", "1. Super Administrator: ")
    add_bullet("Institutional configuration, academic hierarchy setup, staff account provisioning, and white-label landing page CMS customization.", "2. University Administrator: ")
    add_bullet("Student intake, application review, cohort management, semester registration, program transfers, and student holds.", "3. Registrar / Admission Officer: ")
    add_bullet("Fiscal calendar control, period opening/closing, budget approvals, high-value voucher authorization, and financial statement sign-off.", "4. Financial Controller / Manager: ")
    add_bullet("Journal voucher drafting, cheque portfolio management, fixed asset capitalisation, depreciation batch dispatch, and reconciliation.", "5. Senior Accountant: ")
    add_bullet("Front-desk student billing, multi-channel fee receipting, daily cash drawer reconciliation, and receipt re-printing.", "6. Cashier / Teller: ")
    add_bullet("Faculty-level student rosters, enrollment statistics, academic standing reviews, and quota monitoring.", "7. Dean / Academic Head: ")
    add_bullet("Online application tracking, fee invoice review, digital payments, registration proofs, and official statement downloads.", "8. Student / Guardian: ")

    add_callout(
        "The RBAC engine enforces strict Segregation of Duties. The platform refuses any role definition granting: "
        "(a) Voucher Creation AND Approval; (b) Fee Matrix Maintenance AND Student Discount Approval; "
        "(c) Vendor Master Maintenance AND Payment Approval; (d) Journal Entry Posting AND Fiscal Period Unlocking.",
        "REQ-SOD-01: SEGREGATION OF DUTIES"
    )

    # 3. Functional Requirements
    add_h1("3. Specific Functional Requirements")

    # Module 1
    add_h2("Module 1: White-Label Landing Page & CMS Engine")
    add_p("Enables each university client to deploy a branded public web portal without software modifications.")
    add_bullet("Configurable university name (Arabic/English), logo (SVG/PNG), favicon, motto, HSL brand color palette, typography tokens, and social links.", "REQ-LP-01: Dynamic Theme & Branding: ")
    add_bullet("Hero banner with customizable headline, video/image background, and dynamic action buttons ('Apply Online', 'Student Portal', 'Staff Portal').", "REQ-LP-02: Hero & Call-to-Action: ")
    add_bullet("Public interactive catalog of Faculties, Degrees (B.Sc., M.Sc., Diploma), program curricula, admission criteria, and tuition fee schedules.", "REQ-LP-03: Faculties & Programs Explorer: ")
    add_bullet("Multi-step online application wizard with certificate upload, ranking of program choices, tracking number generation, and optional online application fee gateway.", "REQ-LP-04: Online Admissions Portal: ")
    add_bullet("Rich-text publishing system for institutional announcements, academic deadlines, examination calendars, and semester start dates.", "REQ-LP-05: News & Academic Calendar CMS: ")
    add_bullet("Interactive Google Maps embed, campus contact directories, and prospective student inquiry forms.", "REQ-LP-06: Campus Information & Contact: ")

    # Module 2
    add_h2("Module 2: Academic Hierarchy & Configuration Module")
    add_bullet("Management of university academic divisions (Faculties/Colleges, Departments, Research Centers).", "REQ-AC-01: Faculty & Department Management: ")
    add_bullet("Configuration of degree programs, duration (semesters/years), and required credit/course thresholds.", "REQ-AC-02: Academic Programs & Degrees: ")
    add_bullet("Academic calendar configuration (Years e.g., 2026/2027, Semesters: Fall/Spring/Summer, and Student Batches).", "REQ-AC-03: Academic Years & Batch Cycles: ")
    add_bullet("Effective-dated matrix defining Program × Batch × Admission Channel (General/Private/Foreign/Staff) × Nationality Category → Fee Schedule. Every modification generates an immutable version record.", "REQ-AC-04: Comprehensive Fee Matrix Configuration: ")
    add_bullet("Master classification tables for national admission channels, expatriate quotas, and foreign student fee rules.", "REQ-AC-05: Nationalities & Admission Categories: ")
    add_bullet("Hierarchical cost center allocation matching institutional organizational structures (Faculties, Clinics, Labs, Administration).", "REQ-AC-06: Cost Center Master: ")

    # Module 3
    add_h2("Module 3: Student Admission & Digital Profile Management")
    add_bullet("Capture of structured 4-part Arabic name and 4-part English name, unique University ID generation via configurable masks (e.g., [YEAR]-[PROG]-[SEQ]), demographics (DOB Gregorian & Hijri, Nationality, National ID), and guardian contacts.", "REQ-ST-01: Multi-Step Profile Builder: ")
    add_bullet("Standardized health screening recording blood type, Hepatitis B, HIV/AIDS, vaccination history, allergies, medical officer notes, and medical clearance status (Fit / Unfit / Conditional).", "REQ-ST-02: Medical Fitness Records: ")
    add_bullet("Sub-100ms debounced search with mandatory Arabic text normalisation (folding alef variants, taa marbuta/haa, yaa/alef maqsura, and tatweel).", "REQ-ST-03: Advanced Student Search & Filter: ")
    add_bullet("Comprehensive edit history tracking field-level changes with actor ID, timestamp, and before/after values.", "REQ-ST-04: Student Profile Audit History: ")
    add_bullet("Per-program mandatory document verification checklist with document expiry tracking for international students.", "REQ-ST-05: Document Checklist & Verification: ")

    # Module 4
    add_h2("Module 4: Annual / Semester Registration & Student Lifecycle")
    add_bullet("Student lookup → dynamic fee calculation based on active fee matrix version → application of authorized discount scheme → net fee computation.", "REQ-REG-01: Semester Registration Workflow: ")
    add_bullet("Completing registration atomically creates the registration record AND posts a double-entry journal transaction (Dr Student AR, Cr Unearned Tuition Revenue) within the same database transaction.", "REQ-REG-02: Automatic General Ledger Posting: ")
    add_bullet("Cancellation of registration generates a linked reversing journal entry; never edits or deletes historical postings.", "REQ-REG-03: Registration Cancellation & Reversal: ")
    add_bullet("Transfers between programs trigger an automated financial adjustment: reversal of prior program charges and posting of new program fee structure.", "REQ-REG-04: Student Program Transfer: ")
    add_bullet("Generates printable digital registration cards and student ID cards featuring a tamper-proof QR code resolving to a verification API.", "REQ-REG-05: Digital Registration Proof & ID Cards: ")
    add_bullet("System flags (financial, academic, disciplinary, documentary) that automatically restrict registration and downstream services.", "REQ-REG-06: Automated Registration Holds: ")

    # Module 5
    add_h2("Module 5: 5-Tier Chart of Accounts & General Ledger (GL)")
    add_bullet("5-level normalized tree: (1) Major Class, (2) Group, (3) Sub-Group, (4) Parent Account, (5) Detail Account. Only Level-5 accounts accept postings.", "REQ-FIN-01: 5-Tier Chart of Accounts (دليل الحسابات): ")
    add_bullet("Accounts designated as Control Accounts (Student AR, Sponsor AR, Vendor AP) accept postings exclusively via sub-ledgers to prevent ledger divergence.", "REQ-FIN-02: Sub-Ledger Control Accounts: ")
    add_bullet("N-debit to N-credit journal entry grid with real-time balance validation, multi-currency conversion, FX gain/loss handling, and document attachments.", "REQ-FIN-03: Multi-Line General Journal Vouchers: ")
    add_bullet("4-stage workflow (Draft → Pending Review → Pending Approval → Posted). Approvals above configurable thresholds require MFA re-authentication.", "REQ-FIN-04: Maker-Checker Multi-Tier Voucher Approval: ")
    add_bullet("Posted vouchers cannot be modified. Corrections require generating a linked reversal voucher with mandatory audit justification.", "REQ-FIN-05: Audit-Compliant Voucher Reversals: ")
    add_bullet("Per-tenant, per-fiscal-year, per-document-type sequential gapless numbering enforced by transactional database sequences.", "REQ-FIN-06: Sequential Document Numbering: ")

    # Module 6
    add_h2("Module 6: Cashiering, Student Billing & Accounts Receivable")
    add_bullet("Cashier POS interface supporting Cash, Bank Deposit, Cheque, and Online Gateway payments with instant bilingual thermal/A4 receipt generation.", "REQ-CSH-01: Multi-Channel Student Fee Collection: ")
    add_bullet("Configurable installment rules per program (e.g., 50% registration, 25% midterm, 25% final) with automated SMS/Email due date notifications.", "REQ-CSH-02: Installment Plans & Due Dates: ")
    add_bullet("Disbursement vouchers for operational expenses, refunds, and vendor payments with maker-checker approval controls.", "REQ-CSH-03: General Payment Vouchers (سند صرف): ")
    add_bullet("Vouchers for non-student revenues (grants, property leases, donations, certificate fees).", "REQ-CSH-04: General Receipt Vouchers (سند قبض): ")
    add_bullet("Provider-agnostic payment gateway architecture (Webhook callbacks, cryptographic signature validation, automatic reconciliation).", "REQ-CSH-05: Online Payment Gateway Engine: ")
    add_bullet("Same-day receipt voiding by authorized supervisors with automatic linked reversal journal creation.", "REQ-CSH-06: Receipt Cancellation Workflow: ")

    # Module 7
    add_h2("Module 7: Cheque Management & Clearing Pipeline")
    add_bullet("Central registry of post-dated and on-demand cheques (Cheque No, Bank, Branch, Due Date, Drawer Name, Amount, Custody).", "REQ-CHQ-01: Cheque Portfolio & Custody Tracking: ")
    add_bullet("Multi-state lifecycle: Received (تحت التحصيل) → Sent to Bank (برسم التحصيل) → Cleared (محصل: Dr Bank, Cr Cheques Rec.) / Bounced (مرتجع: Dr Student AR, Cr Cheques Rec.) / Cancelled.", "REQ-CHQ-02: Cheque Clearing State Machine: ")
    add_bullet("Bounced cheques automatically reinstate student balance, log bank refusal codes, and apply configurable penalty fees.", "REQ-CHQ-03: Bounced Cheque Automation: ")

    # Module 8
    add_h2("Module 8: Institutional Budgeting & Financial Control")
    add_bullet("Annual budget preparation by Fiscal Year × Account × Cost Center, with monthly distributions and revision versioning.", "REQ-BDG-01: Fiscal Budget Preparation: ")
    add_bullet("Real-time available budget calculation: Available = Allocated - Encumbered - Actual. Hard-stop or warning policies on budget overruns.", "REQ-BDG-02: Real-Time Budgetary Control: ")
    add_bullet("Multi-dimensional reporting comparing Budgeted, Encumbered, Actual, Variance, and Utilization %.", "REQ-BDG-03: Budget vs. Actual Variance Reporting: ")

    # Module 9
    add_h2("Module 9: Fixed Assets Management & Depreciation Engine")
    add_bullet("Cataloging of university assets (Lab equipment, IT, furniture, vehicles, buildings) with QR/Barcode, location, custodian, and cost.", "REQ-AST-01: Fixed Asset Registry: ")
    add_bullet("Automated straight-line depreciation engine generating a persisted schedule per asset from capitalisation date to salvage value.", "REQ-AST-02: Depreciation Engine: ")
    add_bullet("Period-end automated posting (Dr Depreciation Expense, Cr Accumulated Depreciation). Runs are idempotent and period-locked.", "REQ-AST-03: Automated Depreciation Journal Dispatch: ")
    add_bullet("Asset sale, scrap, or revaluation with automatic gain/loss calculation and derecognition journals.", "REQ-AST-04: Asset Disposal & Revaluation: ")

    # Module 10
    add_h2("Module 10: Financial Statements, Reports & Business Intelligence")
    add_bullet("Detailed chronological statement of charges, credits, running balance, and receipt numbers.", "REQ-RPT-01: Student Statement of Account (كشف حساب طالب): ")
    add_bullet("Receivable aging buckets (<30, 31-60, 61-90, >90 days) aged from installment due dates, filterable by faculty/program.", "REQ-RPT-02: Aging of Student Receivables: ")
    add_bullet("Multi-level Trial Balance with Opening, Movement, and Closing columns sourced directly from period-balance aggregators.", "REQ-RPT-03: Trial Balance (ميزان المراجعة): ")
    add_bullet("Standardized institutional Balance Sheet formatted from Level 1 to Level 4 hierarchy.", "REQ-RPT-04: Multi-Level Balance Sheet (الميزانية العمومية): ")
    add_bullet("Operational revenues vs. expenses, net surplus/deficit, filterable by Cost Center with comparative prior-period analysis.", "REQ-RPT-05: Income Statement (قائمة الدخل): ")
    add_bullet("Automated verification report ensuring Sub-Ledger totals match General Ledger Control Account balances exactly.", "REQ-RPT-06: Sub-Ledger Reconciliation Report: ")
    add_bullet("High-speed export to Excel (.xlsx), CSV, and shaped bilingual PDF with official university header and digital stamps.", "REQ-RPT-07: Export & Reporting Engine: ")

    # Module 11
    add_h2("Module 11: Multi-Tenancy, Security & Administration")
    add_bullet("PostgreSQL Row-Level Security (RLS) and application middleware enforcing strict tenant isolation.", "REQ-ADM-01: Multi-Tenant Partitioning: ")
    add_bullet("Granular role-based permissions governing view, create, edit, approve, and reverse actions per module.", "REQ-ADM-02: RBAC Matrix: ")
    add_bullet("Cryptographically hash-chained, append-only audit trail logging user, IP, timestamp, action, and full JSON diff.", "REQ-ADM-03: Immutable Tamper-Evident Audit Log: ")
    add_bullet("Zero-reload language switching, full RTL layout mirroring, dual Gregorian/Hijri calendars, and Arabic monetary amount spelling (تفقيط).", "REQ-ADM-04: Bilingual RTL/LTR Engine: ")
    add_bullet("Tiered SaaS subscription plan enforcement controlling active modules, student capacity, and administrative seats.", "REQ-ADM-05: Tenant Licensing & Entitlements: ")

    # Module 12
    add_h2("Module 12: Fiscal Calendar, Period Control & Opening Balances")
    add_bullet("Configurable fiscal years divided into accounting periods with strict states: Future, Open, Closed, Permanently Closed.", "REQ-PER-01: Fiscal Year & Period Definition: ")
    add_bullet("System rejects any posting dated within a Closed fiscal period. Period close requires passing an automated pre-close reconciliation checklist.", "REQ-PER-02: Posting Period Lock: ")
    add_bullet("Dedicated opening balance module for institutional go-live, requiring exact debit/credit and sub-ledger balancing before live operations commence.", "REQ-PER-03: Opening Balances Module: ")
    add_bullet("Automated closing of revenue and expense accounts into Retained Earnings / Surplus with balance carry-forward.", "REQ-PER-04: Year-End Financial Close: ")

    # Module 13
    add_h2("Module 13: Fee Item Catalog & Revenue Recognition")
    add_bullet("Master catalog of billable fees (Tuition, Registration, Labs, Exams, ID, Hostel, Transport) with assigned GL revenue accounts.", "REQ-FEE-01: Fee Item Catalog: ")
    add_bullet("Tuition fees are initially credited to Unearned Fee Income and recognized into revenue across semester periods via scheduled recognition jobs.", "REQ-FEE-02: Deferred Revenue Recognition: ")
    add_bullet("Automated refund calculation based on withdrawal date schedules, reversing unearned income and raising accounts payable.", "REQ-FEE-03: Withdrawal & Refund Governance: ")
    add_bullet("Maintenance of student credit balances resulting from overpayments or approved refunds, auto-applied to future registrations.", "REQ-FEE-04: Student Credit Balances: ")

    # Module 14
    add_h2("Module 14: Student Status Lifecycle Management")
    add_bullet("Formal state machine: Applicant → Admitted → Active → Deferred / Suspended / Withdrawn / Dismissed / Transferred / Graduated → Alumni.", "REQ-LIF-01: Comprehensive Status Model: ")
    add_bullet("Every status transition requires formal justification, supporting documentation, and authorized workflow sign-off.", "REQ-LIF-02: Transition Governance & Audit: ")
    add_bullet("Streamlined re-admission workflow restoring prior academic history and applying the currently active fee matrix version.", "REQ-LIF-03: Re-Admission Processing: ")

    # Module 15
    add_h2("Module 15: Sponsors & Scholarships")
    add_bullet("Registry of corporate sponsors, government ministries, and embassies with contract parameters and billing addresses.", "REQ-SPN-01: Sponsor Master: ")
    add_bullet("Automated billing split apportioning student charges between student self-pay and sponsor AR accounts.", "REQ-SPN-02: Split Funding Engine: ")
    add_bullet("Consolidated invoicing per sponsor per billing cycle with automated aging and payment matching.", "REQ-SPN-03: Sponsor Invoicing & Reconciliation: ")
    add_bullet("Scholarship program management with allocated budgets, eligibility criteria, approval thresholds, and foregone revenue exposure reports.", "REQ-SPN-04: Scholarship Schemes & Discount Governance: ")

    # Module 16
    add_h2("Module 16: Procurement & Accounts Payable with Encumbrance")
    add_bullet("Vendor directory, bank accounts, tax numbers, payment terms, and assigned AP control accounts.", "REQ-PRC-01: Vendor Master: ")
    add_bullet("Full procure-to-pay workflow: Purchase Requisition → Purchase Order (creates Encumbrance) → Goods Receipt → Vendor Invoice → Payment Voucher.", "REQ-PRC-02: Procure-to-Pay Workflow: ")
    add_bullet("Automated commitment against budget lines upon PO approval, converting to accrued liability on goods receipt.", "REQ-PRC-03: Encumbrance Accounting: ")
    add_bullet("Automated three-way matching among Purchase Order, Goods Receipt, and Vendor Invoice with configurable variance tolerances.", "REQ-PRC-04: Three-Way Match Engine: ")
    add_bullet("Vendor invoice aging and batch payment proposal generator.", "REQ-PRC-05: AP Aging & Payment Run: ")

    # Module 17
    add_h2("Module 17: Admissions Capacity & Committee Workflow")
    add_bullet("Real-time seat quotas per Program × Batch × Admission Channel (General, Private, Foreign, Scholarship) preventing over-enrollment.", "REQ-ADM-CAP-01: Seat Quota Management: ")
    add_bullet("Automated screening of secondary certificate scores, prerequisites, and age/nationality rules.", "REQ-ADM-CAP-02: Automated Eligibility Rules: ")
    add_bullet("Committee evaluation portal with candidate ranking and structured decision recording (Accept, Conditional, Waitlist, Reject).", "REQ-ADM-CAP-03: Committee Review & Scoring: ")
    add_bullet("Automated bilingual offer letter generation with expiration deadlines and online seat deposit collection.", "REQ-ADM-CAP-04: Offers & Deposit Payments: ")
    add_bullet("Intelligent duplicate candidate matching based on national ID, passport, and normalized Arabic name.", "REQ-ADM-CAP-05: Duplicate Candidate Detection: ")
    add_bullet("Bulk import engine for centrally-allocated admissions rosters (Excel/CSV) with pre-commit dry-run validation.", "REQ-ADM-CAP-06: Bulk Intake Roster Import: ")

    # Module 18
    add_h2("Module 18: Academic Records & Course Management (Phase 7-8 Roadmap)")
    add_bullet("Course master with codes, bilingual titles, credit/contact hours, and prerequisite/co-requisite dependency trees.", "REQ-ACD-01: Course Catalog & Curriculum: ")
    add_bullet("Term course section scheduling with instructor allocation, room capacity, and timetable clash detection.", "REQ-ACD-02: Timetable & Section Scheduling: ")
    add_bullet("Student course enrollment with prerequisite validation, credit load limits, and automated financial hold blocking (REQ-REG-06).", "REQ-ACD-03: Course Registration Engine: ")
    add_bullet("Session attendance tracking enforcing regional 75% examination eligibility rules with automated risk alerts.", "REQ-ACD-04: Attendance Tracking & Warnings: ")
    add_bullet("Multi-component grade entry, instructor→department head approval workflow, GPA/cGPA calculation, and academic standing evaluation.", "REQ-ACD-05: Assessment, Grading & GPA: ")
    add_bullet("Exam hall seating allocation, invigilator schedules, results publishing, and fee-bearing re-sit workflows.", "REQ-ACD-06: Examination Management: ")
    add_bullet("Official bilingual academic transcripts with QR verification, degree audit validation, and multi-department graduation clearance.", "REQ-ACD-07: Transcripts & Degree Audit: ")

    # 4. Data Architecture
    add_h1("4. Data Architecture & Database Schema")
    add_h2("4.1 Core Relational Schema")
    add_p(
        "The relational database schema is structured into normalized, tenant-isolated tables utilizing strict foreign keys, "
        "check constraints, and immutable audit structures. The core tables include:"
    )

    schema_headers = ["Domain", "Table Name", "Key Attributes & Purpose"]
    schema_rows = [
        ["Platform", "tenants", "id, slug, name_en, name_ar, logo_url, theme_config, functional_currency, fiscal_year_start_month, plan_id, is_active"],
        ["Security", "users, roles, permissions", "User authentication, role definitions, and granular screen/action permission mapping."],
        ["Security", "audit_log", "id, tenant_id, actor_id, ip, occurred_at, action, resource_type, resource_id, before_json, after_json, prev_hash, hash (Append-Only)"],
        ["Security", "idempotency_keys", "key, tenant_id, endpoint, request_hash, response_json, created_at (Guarantees single execution)"],
        ["Platform", "document_sequences", "tenant_id, fiscal_year_id, doc_type, next_value, prefix, padding (Transactional gapless sequence allocator)"],
        ["Academic", "faculties, academic_programs", "Institutional academic hierarchy, program duration, credit structures, and college divisions."],
        ["Academic", "academic_batches, academic_years", "Academic calendar divisions, cohorts, and semester terms."],
        ["Academic", "fee_items, fee_schedules", "Catalog of billable fees and effective-dated, versioned fee pricing matrices."],
        ["Academic", "seat_quotas", "id, tenant_id, program_id, batch_id, admission_channel, capacity, allocated, confirmed."],
        ["Students", "applications", "Prospective student applications, program ranking choices, review scores, decisions, and deposits."],
        ["Students", "students", "id, tenant_id, student_index, national_id, name_ar_1..4, name_en_1..4, name_ar_normalized, dob, status, program_id"],
        ["Students", "student_status_history", "Historical log of student lifecycle transitions with supporting document references."],
        ["Students", "holds", "id, student_id, hold_type (financial/academic/disciplinary), reason, placed_by, cleared_by."],
        ["Students", "medical_records, student_docs", "Student health clearance evaluations and verified uploaded documents."],
        ["Billing", "semester_registrations", "id, tenant_id, student_id, program_id, academic_year, term_id, discount_pct, status, posted_header_id"],
        ["Billing", "charges, installments", "Itemized semester fee charges and structured payment installment schedules."],
        ["Billing", "revenue_recognition_schedule", "Deferred revenue schedules amortizing unearned income across semester periods."],
        ["Sponsors", "sponsors, sponsorships", "Corporate/embassy sponsor contracts, student coverage caps, and co-funding rules."],
        ["General Ledger", "fiscal_years, fiscal_periods", "Accounting calendar, period lock states (Future, Open, Closed, Permanently Closed)."],
        ["General Ledger", "chart_of_accounts", "5-tier accounting tree (Level 1-5), parent keys, normal balance, postable flag, control account flag."],
        ["General Ledger", "cost_centers", "Multi-dimensional organizational cost allocation units."],
        ["General Ledger", "voucher_drafts, approval_events", "Staged voucher drafts and complete multi-tier approval/rejection audit trail."],
        ["General Ledger", "transaction_headers", "id, tenant_id, fiscal_period_id, voucher_no, voucher_type, doc_date, total_amount, currency, posted_by, reversed"],
        ["General Ledger", "transaction_lines", "id, header_id, line_no, account_id, cost_center_id, subledger_id, debit_amount, credit_amount (numeric 19,4)"],
        ["General Ledger", "account_period_balances", "Incremental period balance rollups powering sub-100ms trial balance and balance sheet queries."],
        ["Treasury", "cheques", "id, tenant_id, cheque_no, bank_name, due_date, drawer_name, amount, status, custody, bounce_reason"],
        ["Assets", "fixed_assets, depr_schedule", "Fixed asset registry and pre-computed straight-line depreciation schedules."],
        ["Budgeting", "budgets, budget_lines", "Annual fiscal budgets per account and cost center with real-time encumbrance tracking."],
        ["Procurement", "vendors, purchase_orders", "Vendor master, purchase requisitions, purchase orders, goods receipts, and 3-way matching."]
    ]
    add_styled_table(schema_headers, schema_rows, [1.4, 2.3, 3.2])

    add_h2("4.2 Database-Enforced Invariants")
    inv_headers = ["Invariant Rule", "Database Enforcement Mechanism"]
    inv_rows = [
        ["Double-Entry Voucher Balance", "Deferred PostgreSQL constraint trigger enforcing SUM(debits) = SUM(credits) in functional currency per voucher."],
        ["Gapless Document Numbering", "Transactional sequence table with UNIQUE (tenant_id, fiscal_year_id, voucher_type, voucher_no)."],
        ["Closed Period Lock", "Database trigger on transaction_headers validating fiscal_periods.status = 'Open' prior to insert."],
        ["Strict Level-5 Postability", "Foreign key check constraint ensuring postings only reference accounts where is_postable = TRUE."],
        ["Sub-Ledger Integrity", "CHECK constraint ensuring control accounts always include a valid subledger_id."],
        ["Ledger Immutability", "PostgreSQL revoked UPDATE/DELETE privileges on transaction_headers and transaction_lines."],
        ["Zero Monetary Precision Loss", "All financial columns declared strictly as numeric(19,4)."],
        ["Multi-Tenant Isolation", "PostgreSQL Row-Level Security (RLS) policies applied to all tenant operational tables."],
        ["Single Reversal Constraint", "UNIQUE index on (reverses_id) WHERE reverses_id IS NOT NULL ensuring vouchers are reversed at most once."],
        ["Tamper-Evident Audit Chain", "Cryptographic SHA-256 hash chaining linking each audit log entry to its predecessor (prev_hash)."]
    ]
    add_styled_table(inv_headers, inv_rows, [2.4, 4.5])

    # 5. Non-Functional Requirements
    add_h1("5. Non-Functional Requirements")
    add_h2("5.1 Performance, Scalability & Responsiveness")
    add_bullet("Initial application payload load < 1.5 seconds; client-side route transitions < 200 milliseconds.", "REQ-NFR-01: Page Load & Transition: ")
    add_bullet("Debounced student and account searches resolve in < 100 ms at scale (> 100,000 students, > 1,000,000 journal lines) utilizing trigram and normalized Arabic indexes.", "REQ-NFR-02: Search & Query Latency: ")
    add_bullet("Sustained throughput of 500+ concurrent cashiering and registration transactions with zero deadlocks and zero sequence collisions.", "REQ-NFR-03: High Concurrency: ")

    add_h2("5.2 Security, Compliance & Data Protection")
    add_bullet("Mandatory TLS 1.3 in transit; AES-256 encryption at rest for databases and document vaults.", "REQ-NFR-04: Cryptographic Standards: ")
    add_bullet("Session tokens stored in Secure/HttpOnly/SameSite cookies; Argon2id password hashing; MFA required for high-value financial actions.", "REQ-NFR-05: Identity & Authentication: ")
    add_bullet("Total prohibition of hard deletes or silent edits on financial data. All corrections require explicit linked reversals.", "REQ-NFR-06: Financial Audit Integrity: ")
    add_bullet("Mandatory idempotency tokens on all state-changing financial APIs, preventing duplicate payments or registrations during network retries.", "REQ-NFR-07: Idempotency Enforcement: ")
    add_bullet("All database access executed via parameterized queries or typed ORM models, eliminating SQL injection vectors.", "REQ-NFR-08: Injection Immunity: ")

    add_h2("5.3 Accessibility, Bilingual & Regional Support")
    add_bullet("Native Arabic RTL support with typography (Cairo / Tajawal / IBM Plex Arabic) and English LTR (Inter); correct Arabic PDF text shaping; algorithmic Arabic monetary spelling (تفقيط).", "REQ-NFR-09: Bilingual & RTL Typography: ")
    add_bullet("Simultaneous display and input support for Gregorian and Hijri calendars across all student-facing and official documents.", "REQ-NFR-10: Dual Calendar Engine: ")
    add_bullet("Fully responsive across desktop workstations (1920x1080, 1440x900), tablets, and POS mobile screens.", "REQ-NFR-11: Responsive Layouts: ")
    add_bullet("shadcn UI / Tailwind CSS design tokens supporting seamless Light and Dark modes and custom university brand palettes.", "REQ-NFR-12: Theming & Tokens: ")
    add_bullet("Compliance with WCAG 2.2 Level AA accessibility standards across all public and student portal surfaces.", "REQ-NFR-13: Accessibility: ")

    add_h2("5.4 Availability, Disaster Recovery & Telemetry")
    add_bullet("Recovery Point Objective (RPO) <= 15 minutes; Recovery Time Objective (RTO) <= 4 hours with automated snapshot drills.", "REQ-NFR-14: Backup & Disaster Recovery: ")
    add_bullet("Structured JSON logging with request correlation IDs, real-time alerting on voucher imbalances or sub-ledger discrepancies.", "REQ-NFR-15: System Observability: ")
    add_bullet("Automated tenant data export in open formats (CSV/JSON/SQL) with configurable institutional retention policies.", "REQ-NFR-16: Data Export & Retention: ")

    # 6. Tenant Onboarding
    add_h1("6. Institutional Tenant Onboarding & Initialization")
    add_p(
        "UniFlow enforces a rigorous, 4-stage onboarding protocol to ensure that new institutional tenants go live "
        "with fully reconciled, mathematically verified operational ledgers:"
    )
    add_bullet("Import from standard institutional CSV templates or seed from the pre-packaged Higher Education COA template. Validates 5-tier depth, normal balances, and control account flags.", "Stage 1 — Chart of Accounts Setup: ")
    add_bullet("Import of faculties, degree programs, batch cohorts, fee items, fee schedules, cost centers, and active student rosters with duplicate checking and dry-run validation.", "Stage 2 — Master Data Initialization: ")
    add_bullet("Entry of opening GL balances and per-student, per-sponsor, and per-vendor sub-ledger balances. Go-live is gated on zero variance between total debits/credits and sub-ledger totals.", "Stage 3 — Opening Balances & Reconciliation Gate: ")
    add_bullet("User provisioning, RBAC role assignment, Segregation of Duties review, opening of the first fiscal year/periods, and cashier drawer allocation.", "Stage 4 — Go-Live & Operational Cutover: ")

    # 7. Future Scope
    add_h1("7. Strategic Product Roadmap & Future Extensions")
    add_p("The following strategic capabilities are specified for planned deployment in subsequent product milestones:")
    future_headers = ["Functional Area", "Target Capability", "Strategic Value & Operational Impact"]
    future_rows = [
        ["Cashier Shift Controls", "Shift open/close with denomination counting, Z-reports, and petty cash imprest management.", "Enhanced teller drawer security and cash handling governance."],
        ["Payroll & HR-Lite", "Staff registry, grade pay scales, monthly payroll batch posting, and end-of-service accruals.", "Unifies university human capital expenditure directly into the GL."],
        ["Offline PWA Terminal", "Offline-tolerant cashier Progressive Web App queueing receipts with local cryptographic signing.", "Guarantees cashier continuity during campus network or power outages."],
        ["Automated Communications", "Multi-channel SMS, WhatsApp, and Email dunning pipeline for payment reminders and notices.", "Accelerates student accounts receivable collection cycles."],
        ["Ancillary Billing", "Hostel bed allocation, transportation route subscriptions, and library fee integration.", "Consolidates campus service revenues into the central student ledger."],
        ["Advanced Analytics & BI", "Cash flow forecasting, student enrollment funnel conversion, and regulatory statistical returns.", "Delivers executive decision support for university boards and deans."],
        ["SaaS Platform Billing", "Automated subscription metering, institutional invoice generation, and credit card / wire processing.", "Automates commercial billing for multi-institution deployments."]
    ]
    add_styled_table(future_headers, future_rows, [1.8, 2.5, 2.6])

    doc.save(target_path)
    print(f"DOCX successfully generated at: {target_path}")

def build_html_and_pdf(html_path, pdf_path):
    html_content = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>UniFlow Enterprise ERP - Software Requirements Specification</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=Cairo:wght@400;600;700&display=swap');
  
  @page {
    size: A4 portrait;
    margin: 18mm 15mm 20mm 15mm;
    @top-right {
      content: "UniFlow Enterprise ERP — SRS v2.0.0";
      font-family: 'Inter', sans-serif;
      font-size: 8pt;
      color: #94a3b8;
    }
    @bottom-right {
      content: "Page " counter(page) " of " counter(pages);
      font-family: 'Inter', sans-serif;
      font-size: 8pt;
      color: #94a3b8;
    }
    @bottom-left {
      content: "ISO/IEC/IEEE 29148 Standard Compliant | Confidential";
      font-family: 'Inter', sans-serif;
      font-size: 8pt;
      color: #94a3b8;
    }
  }

  body {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    color: #1f2937;
    line-height: 1.5;
    font-size: 9.5pt;
    background: #fff;
    margin: 0;
    padding: 0;
  }

  .cover-page {
    page-break-after: always;
    padding: 40px 20px 20px 20px;
    border-bottom: 2px solid #e2e8f0;
  }

  .badge {
    display: inline-block;
    padding: 4px 12px;
    border-radius: 9999px;
    background: #eff6ff;
    color: #2563eb;
    font-size: 8.5pt;
    font-weight: 600;
    letter-spacing: 0.5px;
    text-transform: uppercase;
    margin-bottom: 16px;
    border: 1px solid #bfdbfe;
  }

  h1.doc-title {
    font-size: 26pt;
    font-weight: 800;
    color: #1e3a8a;
    line-height: 1.15;
    margin: 0 0 8px 0;
  }

  h2.doc-subtitle {
    font-size: 14pt;
    font-weight: 500;
    color: #2563eb;
    margin: 0 0 24px 0;
  }

  .meta-grid {
    display: grid;
    grid-template-columns: 1fr 1fr;
    gap: 12px;
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 8px;
    padding: 16px;
    margin-bottom: 24px;
  }

  .meta-item {
    font-size: 9pt;
  }

  .meta-label {
    font-weight: 600;
    color: #64748b;
    font-size: 8pt;
    text-transform: uppercase;
    letter-spacing: 0.5px;
  }

  .meta-value {
    color: #0f172a;
    font-weight: 500;
    margin-top: 2px;
  }

  h1 {
    font-size: 15pt;
    font-weight: 700;
    color: #1e3a8a;
    border-bottom: 1.5px solid #cbd5e1;
    padding-bottom: 4px;
    margin-top: 24px;
    margin-bottom: 12px;
    page-break-after: avoid;
  }

  h2 {
    font-size: 11.5pt;
    font-weight: 600;
    color: #2563eb;
    margin-top: 16px;
    margin-bottom: 8px;
    page-break-after: avoid;
  }

  h3 {
    font-size: 10pt;
    font-weight: 600;
    color: #1e293b;
    margin-top: 12px;
    margin-bottom: 4px;
    page-break-after: avoid;
  }

  p {
    margin: 0 0 8px 0;
    text-align: justify;
  }

  ul {
    margin: 4px 0 10px 18px;
    padding: 0;
  }

  li {
    margin-bottom: 4px;
    line-height: 1.45;
  }

  .req-code {
    font-weight: 600;
    color: #0f172a;
  }

  .callout {
    background: #f0f9ff;
    border-left: 4px solid #0284c7;
    border-radius: 0 6px 6px 0;
    padding: 10px 14px;
    margin: 12px 0;
    font-size: 9pt;
    page-break-inside: avoid;
  }

  .callout-title {
    font-weight: 700;
    color: #0369a1;
    margin-bottom: 4px;
    text-transform: uppercase;
    font-size: 8pt;
    letter-spacing: 0.5px;
  }

  table {
    width: 100%;
    border-collapse: collapse;
    margin: 12px 0;
    font-size: 8.5pt;
    page-break-inside: auto;
  }

  tr {
    page-break-inside: avoid;
    page-break-after: auto;
  }

  th {
    background: #1e3a8a;
    color: #ffffff;
    font-weight: 600;
    text-align: left;
    padding: 7px 10px;
    border: 1px solid #1e3a8a;
  }

  td {
    padding: 6px 10px;
    border: 1px solid #e2e8f0;
    vertical-align: top;
  }

  tr:nth-child(even) td {
    background: #f8fafc;
  }

  .arabic-term {
    font-family: 'Cairo', 'Segoe UI', Tahoma, sans-serif;
    color: #0369a1;
    font-weight: 600;
  }

  .diagram-box {
    background: #f8fafc;
    border: 1px solid #cbd5e1;
    border-radius: 6px;
    padding: 12px;
    margin: 12px 0;
    font-family: 'Inter', monospace;
    font-size: 8pt;
    page-break-inside: avoid;
  }

  .flow-grid {
    display: flex;
    justify-content: space-between;
    align-items: center;
    gap: 8px;
    margin: 10px 0;
  }

  .flow-step {
    flex: 1;
    background: #ffffff;
    border: 1px solid #cbd5e1;
    border-radius: 6px;
    padding: 8px;
    text-align: center;
    font-size: 8pt;
    font-weight: 500;
  }

  .flow-arrow {
    color: #94a3b8;
    font-weight: bold;
    font-size: 11pt;
  }

  .page-break {
    page-break-before: always;
  }
</style>
</head>
<body>

<!-- COVER / HEADER SECTION -->
<div class="cover-page">
  <div class="badge">Enterprise Cloud Specification</div>
  <h1 class="doc-title">Software Requirements Specification</h1>
  <h2 class="doc-subtitle">Multi-Tenant University ERP & White-Label Web Platform</h2>
  
  <div class="meta-grid">
    <div class="meta-item">
      <div class="meta-label">Project Name</div>
      <div class="meta-value">UniFlow Enterprise ERP Web Platform</div>
    </div>
    <div class="meta-item">
      <div class="meta-label">Document Version</div>
      <div class="meta-value">2.0.0 (Cloud-Native Edition)</div>
    </div>
    <div class="meta-item">
      <div class="meta-label">Compliance Standards</div>
      <div class="meta-value">IEEE Std 830-1998 / ISO/IEC/IEEE 29148</div>
    </div>
    <div class="meta-item">
      <div class="meta-label">Target Audience</div>
      <div class="meta-value">University Executives, Technical Architects, Dev Team, QA</div>
    </div>
    <div class="meta-item">
      <div class="meta-label">Security & Multitenancy</div>
      <div class="meta-value">PostgreSQL Row-Level Security (RLS) & Argon2id RBAC</div>
    </div>
    <div class="meta-item">
      <div class="meta-label">Localization</div>
      <div class="meta-value">Bilingual Native (Arabic RTL / English LTR) & Dual Calendar</div>
    </div>
  </div>

  <div class="callout">
    <div class="callout-title">Executive Summary</div>
    UniFlow is an institutional-grade, multi-tenant cloud SaaS ERP engineered specifically for higher education. 
    It unifies public white-label admissions, academic hierarchy setup, comprehensive student lifecycle management, 
    and an immutable, double-entry financial accounting and cashiering engine.
  </div>
</div>

<!-- 1. INTRODUCTION -->
<h1>1. Introduction</h1>
<h2>1.1 Purpose & Vision</h2>
<p>
  This document specifies the software requirements for <strong>UniFlow Enterprise ERP</strong>, a modern, responsive, 
  multi-tenant, cloud-native SaaS web platform. The platform is architected for commercial deployment across universities, 
  colleges, and academies, delivering end-to-end institutional management.
</p>
<p>The platform is anchored upon four core architectural pillars:</p>
<ul>
  <li><span class="req-code">1. White-Label Public Portal & Admissions CMS:</span> Custom university branding, program discovery, and self-service student application portals.</li>
  <li><span class="req-code">2. Student Intake & Lifecycle Management:</span> Multi-step registration, biometric/photo capture, medical clearances, and status transitions.</li>
  <li><span class="req-code">3. Enterprise Double-Entry Accounting:</span> 5-tier Chart of Accounts, maker-checker voucher approvals, fee catalogs, deferred revenue recognition, cashiering POS, cheques, budgeting, fixed assets, and financial statements.</li>
  <li><span class="req-code">4. Modern Bilingual UI Engine:</span> Built on React, TypeScript, Tailwind CSS, and shadcn UI, offering zero-reload Arabic (RTL) and English (LTR) localization.</li>
</ul>

<h2>1.2 Definitions, Acronyms, and Abbreviations</h2>
<table>
  <thead>
    <tr>
      <th style="width: 25%;">Term / Acronym</th>
      <th>Definition & Institutional Context</th>
    </tr>
  </thead>
  <tbody>
    <tr><td><strong>SaaS</strong></td><td>Software as a Service — Multi-tenant cloud-hosted institutional architecture.</td></tr>
    <tr><td><strong>COA</strong></td><td>Chart of Accounts (<span class="arabic-term">دليل الحسابات</span>) — 5-tier standardized accounting tree.</td></tr>
    <tr><td><strong>GL</strong></td><td>General Ledger (<span class="arabic-term">دفتر الأستاذ العام</span>) — Core transactional journal store.</td></tr>
    <tr><td><strong>JV</strong></td><td>Journal Voucher (<span class="arabic-term">قيد اليومية</span>) — Double-entry financial balancing transaction.</td></tr>
    <tr><td><strong>RV / PV</strong></td><td>Receipt Voucher (<span class="arabic-term">سند قبض</span>) / Payment Voucher (<span class="arabic-term">سند صرف</span>).</td></tr>
    <tr><td><strong>AR / AP</strong></td><td>Accounts Receivable (Student Billing) / Accounts Payable (Vendor Billing).</td></tr>
    <tr><td><strong>RBAC</strong></td><td>Role-Based Access Control governing granular operational permissions.</td></tr>
    <tr><td><strong>Maker-Checker</strong></td><td>Separation of duties requiring voucher drafting and approval by distinct actors.</td></tr>
    <tr><td><strong>NBV</strong></td><td>Net Book Value (<span class="arabic-term">القيمة الدفترية للأصل</span>) — Asset cost minus accumulated depreciation.</td></tr>
    <tr><td><strong>Encumbrance</strong></td><td>Budget commitment reserved upon Purchase Order approval prior to disbursement.</td></tr>
    <tr><td><strong>Functional Currency</strong></td><td>The declared primary currency in which a tenant maintains its statutory ledgers.</td></tr>
    <tr><td><strong>Idempotency Key</strong></td><td>Client-supplied token guaranteeing repeat API requests execute exactly once.</td></tr>
    <tr><td><span class="arabic-term">تفقيط (Tafqeet)</span></td><td>Algorithmic conversion of monetary amounts into formal written Arabic words on receipts.</td></tr>
  </tbody>
</table>

<!-- 2. OVERALL DESCRIPTION -->
<h1>2. Overall Description & System Architecture</h1>
<h2>2.1 Web Platform Architecture</h2>
<p>
  UniFlow employs a distributed, high-concurrency web architecture comprising a Next.js/React frontend with Tailwind CSS 
  and shadcn UI, REST/GraphQL API microservices, and a PostgreSQL database layer enforcing tenant isolation via Row-Level Security (RLS). 
  Redis provides low-latency idempotency token tracking and session management.
</p>

<h2>2.2 Core System Capabilities Matrix</h2>
<table>
  <thead>
    <tr>
      <th style="width: 20%;">Capability Area</th>
      <th style="width: 48%;">Web Implementation Details</th>
      <th style="width: 32%;">Security & Architectural Control</th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <td><strong>Authentication & IAM</strong></td>
      <td>Argon2id hashing, secure HTTP-only cookies, MFA for financial approvals, granular RBAC.</td>
      <td>OWASP Top 10 Security Baseline</td>
    </tr>
    <tr>
      <td><strong>White-Label CMS</strong></td>
      <td>Per-tenant dynamic theme tokens, custom domains, HSL color palettes, typography selection.</td>
      <td>Zero-Code Multi-Tenant Engine</td>
    </tr>
    <tr>
      <td><strong>Student Admission</strong></td>
      <td>Multi-step intake wizard, 4-part discrete names, photo capture, document checklist.</td>
      <td>Client & Server Schema Validation</td>
    </tr>
    <tr>
      <td><strong>Registration & Billing</strong></td>
      <td>Atomic semester registration coupled with immediate double-entry GL posting in a single transaction.</td>
      <td>ACID Database Transaction</td>
    </tr>
    <tr>
      <td><strong>Chart of Accounts</strong></td>
      <td>5-tier normalized tree with parent keys, account codes, normal balance flags, and bilingual titles.</td>
      <td>Strict Level-5 Postability</td>
    </tr>
    <tr>
      <td><strong>General Ledger</strong></td>
      <td>Immutable transaction headers and lines with database-enforced debit=credit balancing triggers.</td>
      <td>PostgreSQL Trigger Enforced</td>
    </tr>
    <tr>
      <td><strong>Voucher Approvals</strong></td>
      <td>4-stage workflow: Draft → Pending Review → Pending Approval → Posted. Full history logged.</td>
      <td>Immutable State Transition Log</td>
    </tr>
    <tr>
      <td><strong>Cashiering POS</strong></td>
      <td>Multi-channel fee collection (Cash, Bank, Cheque, Gateway) with instant bilingual thermal/A4 receipts.</td>
      <td>Idempotency Token Guaranteed</td>
    </tr>
    <tr>
      <td><strong>Cheque Management</strong></td>
      <td>Lifecycle state machine (Received → In Collection → Deposited → Cleared / Bounced) with automated GL entries.</td>
      <td>State-Triggered Journaling</td>
    </tr>
    <tr>
      <td><strong>Fixed Assets</strong></td>
      <td>Automated straight-line depreciation engine, persisted schedule per asset, period-locked batch dispatch.</td>
      <td>Idempotent Batch Runner</td>
    </tr>
    <tr>
      <td><strong>Budget & Encumbrance</strong></td>
      <td>Budget lines by fiscal year × account × cost center with real-time PO encumbrance reservation.</td>
      <td>Real-Time Overrun Prevention</td>
    </tr>
    <tr>
      <td><strong>Financial Reports</strong></td>
      <td>Maintained period balance rollups (`account_period_balances`) powering sub-100ms trial balance queries.</td>
      <td>< 100ms Query Latency SLA</td>
    </tr>
  </tbody>
</table>

<h2>2.3 System Invariants & Core Design Principles</h2>
<ul>
  <li><span class="req-code">1. Double-Entry Invariant:</span> Every journal voucher must satisfy Sum(Debits) = Sum(Credits) in the functional currency, enforced by database triggers.</li>
  <li><span class="req-code">2. Monetary Precision:</span> All financial values use <code>numeric(19,4)</code> to eliminate floating-point rounding errors.</li>
  <li><span class="req-code">3. Idempotent Financial Execution:</span> Every state-changing financial request requires an idempotency key, preventing duplicate transactions on network retry.</li>
  <li><span class="req-code">4. Financial Immutability:</span> Posted transactions cannot be edited or deleted. Corrections require explicit, linked reversal vouchers with audit comments.</li>
  <li><span class="req-code">5. Multi-Tenant Isolation:</span> Tenant boundaries are enforced mathematically at the database layer via PostgreSQL Row-Level Security (RLS).</li>
</ul>

<div class="callout">
  <div class="callout-title">REQ-SOD-01: Segregation of Duties Mandate</div>
  The platform strictly prohibits any single user role from combining: 
  (a) Creating and Approving the same voucher; 
  (b) Maintaining fee matrices and Authorizing discounts; 
  (c) Editing vendor banking details and Approving vendor disbursements; 
  (d) Posting journal vouchers and Unlocking closed fiscal periods.
</div>

<!-- 3. SPECIFIC FUNCTIONAL REQUIREMENTS -->
<h1>3. Specific Functional Requirements</h1>

<h2>Module 1: White-Label Landing Page & CMS Engine</h2>
<ul>
  <li><span class="req-code">REQ-LP-01: Dynamic Branding:</span> University Admin configures institutional branding tokens (Bilingual Name, Logo, Favicon, Motto, HSL Color Palette, Typography) rendered dynamically across public pages.</li>
  <li><span class="req-code">REQ-LP-02: Hero Banner & CTA:</span> Customizable hero section with rich media backgrounds and direct portal action buttons.</li>
  <li><span class="req-code">REQ-LP-03: Program Explorer:</span> Public searchable catalog of Faculties, Degrees, study plans, career prospects, and published tuition schedules.</li>
  <li><span class="req-code">REQ-LP-04: Online Admissions:</span> Prospective student application wizard with certificate upload, choice ranking, tracking number generation, and online fee payment.</li>
  <li><span class="req-code">REQ-LP-05: News & Calendar:</span> CMS publishing academic calendars, semester milestones, and campus announcements.</li>
  <li><span class="req-code">REQ-LP-06: Campus Contact:</span> Map embed, campus directory, and inquiry routing forms.</li>
</ul>

<h2>Module 2: Academic Hierarchy & Configuration</h2>
<ul>
  <li><span class="req-code">REQ-AC-01: Faculties & Departments:</span> Organizational structure setup for university faculties, academic departments, and research centers.</li>
  <li><span class="req-code">REQ-AC-02: Degree Programs:</span> Degree level definitions (B.Sc., M.Sc., Ph.D., Diploma) with standard durations and credit requirements.</li>
  <li><span class="req-code">REQ-AC-03: Academic Calendar:</span> Academic years, semester terms (Fall, Spring, Summer), and batch cohort definitions.</li>
  <li><span class="req-code">REQ-AC-04: Versioned Fee Matrix:</span> Effective-dated fee matrix by Program × Batch × Admission Type × Nationality Category. Modifications generate an immutable version record.</li>
  <li><span class="req-code">REQ-AC-05: Admission Channels:</span> Master tables for national admission categories, expatriate quotas, and international fee rules.</li>
  <li><span class="req-code">REQ-AC-06: Cost Centers:</span> Multi-tier cost center catalog aligned with faculties and operational divisions.</li>
</ul>

<h2>Module 3: Student Admission & Digital Profile Management</h2>
<ul>
  <li><span class="req-code">REQ-ST-01: Profile Builder:</span> Structured 4-part Arabic and English names, unique University ID generation via configurable masks, demographics (DOB Gregorian/Hijri, National ID), and guardian profiles.</li>
  <li><span class="req-code">REQ-ST-02: Medical Fitness:</span> Health screening records capturing blood group, infectious disease checks, vaccination history, allergies, and doctor clearance verdicts (Fit / Unfit / Conditional).</li>
  <li><span class="req-code">REQ-ST-03: Normalized Search:</span> Sub-100ms debounced search with Arabic text normalisation (folding alef variants, taa marbuta, yaa/alef maqsura, and tatweel).</li>
  <li><span class="req-code">REQ-ST-04: Profile Audit History:</span> Field-level audit logging of all profile updates with actor ID, timestamp, and before/after values.</li>
  <li><span class="req-code">REQ-ST-05: Document Checklist:</span> Verification checklist for required admission credentials with expiry alerts for renewable documents.</li>
</ul>

<h2>Module 4: Semester Registration & Student Lifecycle</h2>
<ul>
  <li><span class="req-code">REQ-REG-01: Registration Workflow:</span> Student lookup → fee matrix resolution → discount/scholarship application → net tuition calculation.</li>
  <li><span class="req-code">REQ-REG-02: Atomic GL Posting:</span> Registration creation atomically posts a balanced double-entry transaction (Dr Student AR, Cr Unearned Fee Income) in the same database transaction.</li>
  <li><span class="req-code">REQ-REG-03: Cancellation & Reversal:</span> Registration cancellation posts a linked reversal voucher, preserving historical records.</li>
  <li><span class="req-code">REQ-REG-04: Program Transfers:</span> Program transfer workflow reversing prior program charges and posting new curriculum billing.</li>
  <li><span class="req-code">REQ-REG-05: Digital Registration Cards:</span> Generation of registration proof and student ID cards featuring a verifiable QR code.</li>
  <li><span class="req-code">REQ-REG-06: Registration Holds:</span> Automated holds (financial, academic, disciplinary) restricting registration until cleared.</li>
</ul>

<h2>Module 5: 5-Tier Chart of Accounts & General Ledger</h2>
<ul>
  <li><span class="req-code">REQ-FIN-01: 5-Tier Chart of Accounts:</span> (1) Major Class, (2) Group, (3) Sub-Group, (4) Parent Account, (5) Detail Account. Only Level-5 accounts accept direct postings.</li>
  <li><span class="req-code">REQ-FIN-02: Control Accounts:</span> Accounts designated as Student AR, Sponsor AR, or Vendor AP accept postings exclusively through sub-ledgers.</li>
  <li><span class="req-code">REQ-FIN-03: Multi-Line Journal Vouchers:</span> N-debit to N-credit journal entry grid with real-time balancing, multi-currency conversion, FX gain/loss handling, and file attachments.</li>
  <li><span class="req-code">REQ-FIN-04: Maker-Checker Approval:</span> 4-stage workflow: Draft → Pending Review → Pending Approval → Posted. High-value approvals require MFA re-authentication.</li>
  <li><span class="req-code">REQ-FIN-05: Immutable Voucher Reversals:</span> Posted vouchers cannot be modified. Adjustments require creating a linked reversal voucher with mandatory justification.</li>
  <li><span class="req-code">REQ-FIN-06: Document Sequences:</span> Per-tenant, per-fiscal-year gapless document numbering managed via transactional sequence allocators.</li>
</ul>

<h2>Module 6: Cashiering POS & Accounts Receivable</h2>
<ul>
  <li><span class="req-code">REQ-CSH-01: Multi-Channel Receipts:</span> Cashier terminal supporting Cash, Bank Transfer, Cheque, and Online Gateway payments with instant bilingual thermal/A4 receipts.</li>
  <li><span class="req-code">REQ-CSH-02: Installment Plans:</span> Configurable installment schedules per program (e.g. 50% registration, 25% midterm, 25% final) with automated SMS/Email reminders.</li>
  <li><span class="req-code">REQ-CSH-03: Payment Vouchers:</span> General disbursement vouchers for vendor invoices, refunds, and operating expenses with maker-checker approvals.</li>
  <li><span class="req-code">REQ-CSH-04: General Receipts:</span> Receipt vouchers for non-student revenues (grants, leases, donations, certificate fees).</li>
  <li><span class="req-code">REQ-CSH-05: Payment Gateway Engine:</span> Provider-agnostic payment gateway interface with webhook signature verification and automated reconciliation.</li>
  <li><span class="req-code">REQ-CSH-06: Receipt Cancellation:</span> Same-day receipt voiding by authorized supervisors generating automated linked reversal journals.</li>
</ul>

<h2>Module 7: Cheque Management & Clearing Pipeline</h2>
<ul>
  <li><span class="req-code">REQ-CHQ-01: Portfolio Tracking:</span> Registry of post-dated and on-demand cheques capturing number, drawee bank, due date, drawer name, and custody status.</li>
  <li><span class="req-code">REQ-CHQ-02: Clearing State Machine:</span> Transition pipeline: Received → Sent to Bank → Cleared (Dr Bank, Cr Cheques Rec.) / Bounced (Dr Student AR, Cr Cheques Rec.) / Cancelled.</li>
  <li><span class="req-code">REQ-CHQ-03: Bounced Cheque Handling:</span> Bounces reinstate student receivable balance, record bank refusal codes, and apply configurable penalty fees.</li>
</ul>

<h2>Module 8: Institutional Budgeting & Financial Control</h2>
<ul>
  <li><span class="req-code">REQ-BDG-01: Annual Fiscal Budget:</span> Budgets configured by Fiscal Year × Account × Cost Center with monthly/quarterly allocation curves.</li>
  <li><span class="req-code">REQ-BDG-02: Real-Time Budgetary Control:</span> Available = Allocated - Encumbered - Actual. System enforces configurable warning or hard-stop policies on overruns.</li>
  <li><span class="req-code">REQ-BDG-03: Variance Reporting:</span> Multi-dimensional reporting comparing Budgeted, Encumbered, Actual, Available, and Utilization %.</li>
</ul>

<h2>Module 9: Fixed Assets & Depreciation Engine</h2>
<ul>
  <li><span class="req-code">REQ-AST-01: Asset Registry:</span> Asset tracking (Lab equipment, IT, furniture, vehicles, buildings) with QR/Barcode, location, custodian, and historical cost.</li>
  <li><span class="req-code">REQ-AST-02: Depreciation Engine:</span> Automated straight-line depreciation engine generating persisted period schedules per asset.</li>
  <li><span class="req-code">REQ-AST-03: Automated Batch Journals:</span> Period-end automated batch posting (Dr Depreciation Expense, Cr Accumulated Depreciation). Runs are idempotent and period-locked.</li>
  <li><span class="req-code">REQ-AST-04: Asset Disposal & Revaluation:</span> Asset retirement, scrap, or revaluation with automated gain/loss calculation.</li>
</ul>

<h2>Module 10: Financial Statements & Business Intelligence</h2>
<ul>
  <li><span class="req-code">REQ-RPT-01: Student Statement:</span> Chronological ledger statement of charges, payments, and running balance with receipt references.</li>
  <li><span class="req-code">REQ-RPT-02: Receivable Aging:</span> Aging buckets (&lt;30, 31-60, 61-90, &gt;90 days) aged from installment due dates, filterable by faculty/program.</li>
  <li><span class="req-code">REQ-RPT-03: Trial Balance:</span> Multi-level Trial Balance with Opening, Movement, and Closing columns sourced directly from period-balance aggregators.</li>
  <li><span class="req-code">REQ-RPT-04: Multi-Level Balance Sheet:</span> Institutional Balance Sheet structured from Level 1 summary to Level 4 detail.</li>
  <li><span class="req-code">REQ-RPT-05: Income Statement:</span> Operating revenues vs. expenses, net surplus/deficit, filterable by Cost Center with comparative prior-period analysis.</li>
  <li><span class="req-code">REQ-RPT-06: Sub-Ledger Reconciliation:</span> Automated verification ensuring Sub-Ledger totals match General Ledger Control Account balances exactly.</li>
  <li><span class="req-code">REQ-RPT-07: Export Engine:</span> Export to Excel (.xlsx), CSV, and shaped bilingual PDF with official university header and digital stamps.</li>
</ul>

<h2>Module 11: Multi-Tenancy & Platform Administration</h2>
<ul>
  <li><span class="req-code">REQ-ADM-01: Multi-Tenant Architecture:</span> PostgreSQL Row-Level Security (RLS) and application middleware enforcing strict tenant isolation.</li>
  <li><span class="req-code">REQ-ADM-02: Granular RBAC:</span> Configurable role permissions governing view, create, edit, approve, and reverse actions per module.</li>
  <li><span class="req-code">REQ-ADM-03: Tamper-Evident Audit Log:</span> Cryptographically hash-chained, append-only audit trail logging user, IP, timestamp, action, and JSON diffs.</li>
  <li><span class="req-code">REQ-ADM-04: Bilingual RTL/LTR Engine:</span> Zero-reload language switching, native RTL mirroring, dual Gregorian/Hijri calendars, and Arabic monetary amount spelling (<span class="arabic-term">تفقيط</span>).</li>
  <li><span class="req-code">REQ-ADM-05: Subscription Licensing:</span> SaaS tier enforcement controlling active modules, student capacity, and administrative seats.</li>
</ul>

<h2>Module 12: Fiscal Calendar & Opening Balances</h2>
<ul>
  <li><span class="req-code">REQ-PER-01: Fiscal Periods:</span> Fiscal years divided into monthly periods with states: Future, Open, Closed, Permanently Closed.</li>
  <li><span class="req-code">REQ-PER-02: Posting Period Lock:</span> System rejects postings into Closed periods. Period close requires passing automated reconciliation checks.</li>
  <li><span class="req-code">REQ-PER-03: Opening Balances:</span> Structured opening balance entry requiring exact debit=credit and sub-ledger reconciliation prior to go-live.</li>
  <li><span class="req-code">REQ-PER-04: Year-End Close:</span> Closing revenue and expense accounts into Retained Earnings / Surplus with balance carry-forward.</li>
</ul>

<h2>Module 13: Fee Item Catalog & Revenue Recognition</h2>
<ul>
  <li><span class="req-code">REQ-FEE-01: Fee Catalog:</span> Catalog of billable fees (Tuition, Registration, Labs, Exams, ID, Hostel, Transport) with designated GL revenue accounts.</li>
  <li><span class="req-code">REQ-FEE-02: Deferred Revenue:</span> Tuition billing credits Unearned Fee Income, recognized into revenue across semester periods via scheduled recognition jobs.</li>
  <li><span class="req-code">REQ-FEE-03: Withdrawal Refunds:</span> Automated refund calculation based on withdrawal date schedules, reversing unearned income and raising accounts payable.</li>
  <li><span class="req-code">REQ-FEE-04: Credit Balances:</span> Maintenance of student credit balances resulting from overpayments, auto-applied to future registrations.</li>
</ul>

<h2>Module 14: Student Status Lifecycle</h2>
<ul>
  <li><span class="req-code">REQ-LIF-01: Status Model:</span> State machine: Applicant → Admitted → Active → Deferred / Suspended / Withdrawn / Dismissed / Transferred / Graduated → Alumni.</li>
  <li><span class="req-code">REQ-LIF-02: Transition Governance:</span> Status changes require recorded justification, supporting documentation, and authorized workflow sign-off.</li>
  <li><span class="req-code">REQ-LIF-03: Re-Admission:</span> Re-admission workflow restoring prior academic history and applying the currently active fee matrix.</li>
</ul>

<h2>Module 15: Sponsors & Scholarships</h2>
<ul>
  <li><span class="req-code">REQ-SPN-01: Sponsor Master:</span> Registry of corporate sponsors, government ministries, and embassies with contract parameters and billing addresses.</li>
  <li><span class="req-code">REQ-SPN-02: Split Funding:</span> Automated billing split apportioning student charges between student self-pay and sponsor AR accounts.</li>
  <li><span class="req-code">REQ-SPN-03: Sponsor Invoicing:</span> Consolidated billing per sponsor per cycle with automated aging and payment matching.</li>
  <li><span class="req-code">REQ-SPN-04: Scholarship Governance:</span> Scholarship programs with allocated budgets, eligibility rules, approval thresholds, and foregone revenue exposure reports.</li>
</ul>

<h2>Module 16: Procurement & Accounts Payable with Encumbrance</h2>
<ul>
  <li><span class="req-code">REQ-PRC-01: Vendor Master:</span> Vendor directory, bank accounts, tax numbers, payment terms, and assigned AP control accounts.</li>
  <li><span class="req-code">REQ-PRC-02: Procure-to-Pay Workflow:</span> Requisition → Purchase Order (creates Encumbrance) → Goods Receipt → Vendor Invoice → Payment Voucher.</li>
  <li><span class="req-code">REQ-PRC-03: Encumbrance Accounting:</span> Automated commitment against budget lines upon PO approval, converting to accrued liability on goods receipt.</li>
  <li><span class="req-code">REQ-PRC-04: Three-Way Match:</span> Automated three-way matching among Purchase Order, Goods Receipt, and Vendor Invoice within configurable tolerances.</li>
  <li><span class="req-code">REQ-PRC-05: AP Aging & Payment Run:</span> Vendor invoice aging and batch payment proposal generator.</li>
</ul>

<h2>Module 17: Admissions Capacity & Committee Workflow</h2>
<ul>
  <li><span class="req-code">REQ-ADM-CAP-01: Seat Quota Management:</span> Live seat quotas per Program × Batch × Admission Channel preventing over-enrollment.</li>
  <li><span class="req-code">REQ-ADM-CAP-02: Automated Eligibility:</span> Screening of secondary certificate scores, required subjects, and age/nationality constraints.</li>
  <li><span class="req-code">REQ-ADM-CAP-03: Committee Review:</span> Committee scoring portal with candidate ranking and structured decision recording (Accept, Conditional, Waitlist, Reject).</li>
  <li><span class="req-code">REQ-ADM-CAP-04: Offers & Deposits:</span> Automated bilingual offer letter generation with expiration deadlines and online seat deposit collection.</li>
  <li><span class="req-code">REQ-ADM-CAP-05: Duplicate Detection:</span> Candidate matching on national ID, passport, and normalized Arabic name to prevent duplicate applications.</li>
  <li><span class="req-code">REQ-ADM-CAP-06: Bulk Intake Import:</span> Bulk import engine for centrally-allocated admissions rosters with pre-commit dry-run validation.</li>
</ul>

<h2>Module 18: Academic Records & Course Management (Phase 7-8 Roadmap)</h2>
<ul>
  <li><span class="req-code">REQ-ACD-01: Course Catalog:</span> Course master with codes, bilingual titles, credit/contact hours, and prerequisite/co-requisite trees.</li>
  <li><span class="req-code">REQ-ACD-02: Timetable Scheduling:</span> Section scheduling with instructor assignment, room capacity, and timetable clash detection.</li>
  <li><span class="req-code">REQ-ACD-03: Course Registration:</span> Student course enrollment with prerequisite verification, credit load limits, and automated financial hold blocking.</li>
  <li><span class="req-code">REQ-ACD-04: Attendance Tracking:</span> Session attendance tracking enforcing regional 75% examination eligibility rules with automated risk alerts.</li>
  <li><span class="req-code">REQ-ACD-05: Assessment & GPA:</span> Multi-component grade entry, approval workflow, GPA/cGPA calculation, and academic standing evaluation.</li>
  <li><span class="req-code">REQ-ACD-06: Examination Management:</span> Exam hall seating allocation, invigilator schedules, results publishing, and fee-bearing re-sit workflows.</li>
  <li><span class="req-code">REQ-ACD-07: Transcripts & Degree Audit:</span> Official bilingual transcripts with QR verification, degree audit validation, and graduation clearance.</li>
</ul>

<!-- 4. DATA ARCHITECTURE -->
<h1>4. Data Architecture & Database Schema</h1>
<h2>4.1 Database-Enforced Invariants</h2>
<table>
  <thead>
    <tr>
      <th style="width: 35%;">Invariant Rule</th>
      <th>Database Enforcement Mechanism</th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <td><strong>Double-Entry Balance</strong></td>
      <td>Deferred PostgreSQL constraint trigger enforcing SUM(debits) = SUM(credits) in functional currency per voucher header.</td>
    </tr>
    <tr>
      <td><strong>Gapless Document Numbering</strong></td>
      <td>Transactional sequence table with UNIQUE (tenant_id, fiscal_year_id, voucher_type, voucher_no).</td>
    </tr>
    <tr>
      <td><strong>Closed Period Lock</strong></td>
      <td>Database trigger on transaction_headers validating fiscal_periods.status = 'Open' prior to insert.</td>
    </tr>
    <tr>
      <td><strong>Strict Level-5 Postability</strong></td>
      <td>Foreign key check constraint ensuring postings only reference accounts where is_postable = TRUE.</td>
    </tr>
    <tr>
      <td><strong>Sub-Ledger Control Integrity</strong></td>
      <td>CHECK constraint ensuring control accounts always include a valid subledger_id.</td>
    </tr>
    <tr>
      <td><strong>Ledger Immutability</strong></td>
      <td>PostgreSQL revoked UPDATE/DELETE privileges on transaction_headers and transaction_lines.</td>
    </tr>
    <tr>
      <td><strong>Zero Precision Loss</strong></td>
      <td>All financial columns declared strictly as numeric(19,4).</td>
    </tr>
    <tr>
      <td><strong>Multi-Tenant Isolation</strong></td>
      <td>PostgreSQL Row-Level Security (RLS) policies applied to all tenant operational tables.</td>
    </tr>
    <tr>
      <td><strong>Single Reversal Guarantee</strong></td>
      <td>UNIQUE index on (reverses_id) WHERE reverses_id IS NOT NULL ensuring vouchers are reversed at most once.</td>
    </tr>
    <tr>
      <td><strong>Tamper-Evident Audit Trail</strong></td>
      <td>Cryptographic SHA-256 hash chaining linking each audit log entry to its predecessor (prev_hash).</td>
    </tr>
  </tbody>
</table>

<!-- 5. NON-FUNCTIONAL REQUIREMENTS -->
<h1>5. Non-Functional Requirements</h1>
<h2>5.1 Performance & Concurrency</h2>
<ul>
  <li><span class="req-code">REQ-NFR-01: Page Load:</span> Initial application load &lt; 1.5 seconds; client-side route transitions &lt; 200 milliseconds.</li>
  <li><span class="req-code">REQ-NFR-02: Query Latency:</span> Debounced student and account searches resolve in &lt; 100 ms at scale (&gt; 100k students, &gt; 1M journal lines).</li>
  <li><span class="req-code">REQ-NFR-03: High Concurrency:</span> Sustained throughput of 500+ concurrent cashiering and registration transactions with zero deadlocks.</li>
</ul>

<h2>5.2 Security & Integrity</h2>
<ul>
  <li><span class="req-code">REQ-NFR-04: Cryptographic Standards:</span> TLS 1.3 in transit; AES-256 encryption at rest for databases and document vaults.</li>
  <li><span class="req-code">REQ-NFR-05: Authentication:</span> Secure/HttpOnly/SameSite session cookies; Argon2id password hashing; MFA for financial approvals.</li>
  <li><span class="req-code">REQ-NFR-06: Financial Immutability:</span> Prohibition of hard deletes or silent edits. All corrections require explicit linked reversals.</li>
  <li><span class="req-code">REQ-NFR-07: Financial Idempotency:</span> Mandatory idempotency tokens on all state-changing financial APIs.</li>
  <li><span class="req-code">REQ-NFR-08: Injection Immunity:</span> Parameterized prepared statements and typed ORMs across all data access paths.</li>
</ul>

<h2>5.3 Regional Localization & UI Ergonomics</h2>
<ul>
  <li><span class="req-code">REQ-NFR-09: Bilingual RTL/LTR:</span> Native Arabic RTL with Cairo/Inter typography; correct Arabic PDF text shaping; algorithmic Arabic monetary spelling (<span class="arabic-term">تفقيط</span>).</li>
  <li><span class="req-code">REQ-NFR-10: Dual Calendar:</span> Simultaneous Gregorian and Hijri calendar support across all student-facing and official documents.</li>
  <li><span class="req-code">REQ-NFR-11: Responsive Layouts:</span> Fluid layouts across desktop displays, tablets, and POS mobile screens.</li>
  <li><span class="req-code">REQ-NFR-12: Theming Tokens:</span> shadcn UI / Tailwind CSS tokens supporting Light/Dark modes and university brand palettes.</li>
  <li><span class="req-code">REQ-NFR-13: Accessibility:</span> WCAG 2.2 Level AA compliance across all public and student portal surfaces.</li>
</ul>

<h2>5.4 Availability & Operability</h2>
<ul>
  <li><span class="req-code">REQ-NFR-14: Backup & Recovery:</span> RPO &le; 15 minutes; RTO &le; 4 hours with automated restore drill verification.</li>
  <li><span class="req-code">REQ-NFR-15: Observability:</span> Structured JSON logging with request correlation IDs and real-time alerting on ledger variances.</li>
  <li><span class="req-code">REQ-NFR-16: Data Export:</span> Automated tenant data export in open formats (CSV/JSON/SQL) with configurable retention policies.</li>
</ul>

<!-- 6. TENANT ONBOARDING -->
<h1>6. Institutional Tenant Onboarding & Initialization</h1>
<p>
  UniFlow enforces a 4-stage onboarding protocol ensuring that institutional tenants go live 
  with fully reconciled, mathematically verified operational ledgers:
</p>
<ul>
  <li><span class="req-code">Stage 1 — Chart of Accounts Setup:</span> Import from institutional CSV templates or seed from the pre-packaged Higher Education COA template. Validates 5-tier depth, normal balances, and control account flags.</li>
  <li><span class="req-code">Stage 2 — Master Data Initialization:</span> Import of faculties, degree programs, batch cohorts, fee items, fee schedules, cost centers, and active student rosters with duplicate checking.</li>
  <li><span class="req-code">Stage 3 — Opening Balances & Reconciliation Gate:</span> Entry of opening GL balances and per-student, per-sponsor, and per-vendor sub-ledger balances. Go-live is gated on zero variance between total debits/credits and sub-ledger totals.</li>
  <li><span class="req-code">Stage 4 — Go-Live & Operational Cutover:</span> User provisioning, RBAC role assignment, Segregation of Duties review, opening of the first fiscal year/periods, and cashier drawer allocation.</li>
</ul>

<!-- 7. FUTURE ROADMAP -->
<h1>7. Strategic Product Roadmap & Future Extensions</h1>
<table>
  <thead>
    <tr>
      <th style="width: 22%;">Functional Area</th>
      <th style="width: 38%;">Target Capability</th>
      <th style="width: 40%;">Strategic Value & Impact</th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <td><strong>Cashier Shift Controls</strong></td>
      <td>Shift open/close with denomination counting, Z-reports, and petty cash imprest management.</td>
      <td>Enhanced teller drawer security and cash handling governance.</td>
    </tr>
    <tr>
      <td><strong>Payroll & HR-Lite</strong></td>
      <td>Staff registry, grade pay scales, monthly payroll batch posting, and end-of-service accruals.</td>
      <td>Unifies university human capital expenditure directly into the GL.</td>
    </tr>
    <tr>
      <td><strong>Offline PWA Terminal</strong></td>
      <td>Offline-tolerant cashier Progressive Web App queueing receipts with local cryptographic signing.</td>
      <td>Guarantees cashier continuity during campus network or power outages.</td>
    </tr>
    <tr>
      <td><strong>Automated Communications</strong></td>
      <td>Multi-channel SMS, WhatsApp, and Email dunning pipeline for payment reminders and notices.</td>
      <td>Accelerates student accounts receivable collection cycles.</td>
    </tr>
    <tr>
      <td><strong>Ancillary Billing</strong></td>
      <td>Hostel bed allocation, transportation route subscriptions, and library fee integration.</td>
      <td>Consolidates campus service revenues into the central student ledger.</td>
    </tr>
    <tr>
      <td><strong>Advanced Analytics & BI</strong></td>
      <td>Cash flow forecasting, student enrollment funnel conversion, and regulatory statistical returns.</td>
      <td>Delivers executive decision support for university boards and deans.</td>
    </tr>
    <tr>
      <td><strong>SaaS Platform Billing</strong></td>
      <td>Automated subscription metering, institutional invoice generation, and credit card / wire processing.</td>
      <td>Automates commercial billing for multi-institution deployments.</td>
    </tr>
  </tbody>
</table>

<p style="margin-top: 30px; text-align: center; color: #94a3b8; font-size: 8pt;">
  — End of Software Requirements Specification (UniFlow Enterprise ERP v2.0.0) —
</p>

</body>
</html>
"""

    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"HTML successfully generated at: {html_path}")

    # Compile HTML to PDF using Chrome Headless
    chrome_path = r"C:\Program Files\Google\Chrome\Application\chrome.exe"
    if not os.path.exists(chrome_path):
        chrome_path = r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
    
    cmd = [
        chrome_path,
        "--headless=new",
        "--disable-gpu",
        "--no-pdf-header-footer",
        f"--print-to-pdf={pdf_path}",
        html_path
    ]
    
    print(f"Running Chrome headless command: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode == 0:
        print(f"PDF successfully generated at: {pdf_path}")
    else:
        print(f"Error generating PDF: {result.stderr}")

if __name__ == "__main__":
    base_dir = r"e:\coding\dr. Ayman work\E-Unvevircity"
    docx_path = os.path.join(base_dir, "srs_university_erp.docx")
    html_path = os.path.join(base_dir, "srs_university_erp.html")
    pdf_path = os.path.join(base_dir, "srs_university_erp.pdf")

    print("Building DOCX...")
    build_docx(docx_path)

    print("Building HTML and PDF...")
    build_html_and_pdf(html_path, pdf_path)
    print("All tasks completed successfully!")
